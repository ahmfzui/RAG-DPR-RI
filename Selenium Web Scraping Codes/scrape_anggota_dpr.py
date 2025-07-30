"""
DPR RI Selenium Scraper - Simple Sentence Format for RAG
Scrapes member data and outputs to simple sentence format optimized for RAG systems
"""

from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.chrome.options import Options
from selenium.common.exceptions import TimeoutException, NoSuchElementException
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
import time
import re
from datetime import datetime
from collections import Counter

class DPRSeleniumSimpleSentences:
    def __init__(self, headless=False):
        self.setup_driver(headless)
        self.members_data = []
        self.base_url = "https://www.dpr.go.id/tentang-dpr/informasi-anggota-dewan"
        
    def setup_driver(self, headless=False):
        """Setup Chrome WebDriver"""
        print("🤖 Setting up Chrome WebDriver...")
        
        chrome_options = Options()
        
        if headless:
            chrome_options.add_argument("--headless")
            
        # Optimization options
        chrome_options.add_argument("--no-sandbox")
        chrome_options.add_argument("--disable-dev-shm-usage")
        chrome_options.add_argument("--disable-blink-features=AutomationControlled")
        chrome_options.add_experimental_option("excludeSwitches", ["enable-automation"])
        chrome_options.add_experimental_option('useAutomationExtension', False)
        
        # Performance options
        chrome_options.add_argument("--disable-images")
        chrome_options.add_argument("--disable-javascript")
        
        # User agent
        chrome_options.add_argument("--user-agent=Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/137.0.0.0 Safari/537.36")
        
        try:
            self.driver = webdriver.Chrome(options=chrome_options)
            self.driver.execute_script("Object.defineProperty(navigator, 'webdriver', {get: () => undefined})")
            self.wait = WebDriverWait(self.driver, 30)
            print("✅ WebDriver ready")
        except Exception as e:
            print(f"❌ Failed to setup WebDriver: {e}")
            print("💡 Make sure ChromeDriver is installed")
            raise
    
    def navigate_to_page(self, page_num=1):
        """Navigate to specific page"""
        if page_num == 1:
            url = self.base_url
        else:
            url = f"{self.base_url}?page={page_num}"
        
        print(f"🌐 Navigating to page {page_num}: {url}")
        
        try:
            self.driver.get(url)
            
            # Wait for member cards to load
            print("⏳ Waiting for member cards to load...")
            
            # Wait for the grid container with member cards
            grid_container = self.wait.until(
                EC.presence_of_element_located((By.CSS_SELECTOR, ".grid.grid-cols-1.gap-6"))
            )
            
            print("✅ Member cards container loaded")
            
            # Additional wait for actual cards
            time.sleep(3)
            
            return True
            
        except TimeoutException:
            print(f"❌ Timeout loading page {page_num}")
            return False
        except Exception as e:
            print(f"❌ Error navigating to page {page_num}: {e}")
            return False
    
    def extract_members_from_current_page(self):
        """Extract all members from current page based on exact HTML structure"""
        print("🔍 Extracting members from current page...")
        
        members = []
        
        try:
            # Find all member cards based on the exact structure
            member_cards = self.driver.find_elements(
                By.CSS_SELECTOR, 
                ".grid.grid-cols-1.gap-6 > .hover\\:cursor-pointer"
            )
            
            print(f"📱 Found {len(member_cards)} member cards")
            
            for i, card in enumerate(member_cards, 1):
                try:
                    member = self.extract_member_from_card(card, i)
                    if member and member.get('nama'):
                        members.append(member)
                        print(f"   ✅ {i:2d}. {member['nama']} - {member['provinsi']}")
                    else:
                        print(f"   ❌ {i:2d}. Failed to extract member data")
                        
                except Exception as e:
                    print(f"   ❌ {i:2d}. Error extracting member: {e}")
                    continue
            
            print(f"✅ Successfully extracted {len(members)} members from page")
            return members
            
        except Exception as e:
            print(f"❌ Error extracting members: {e}")
            return []
    
    def extract_member_from_card(self, card_element, card_number):
        """Extract member data from individual card based on exact HTML structure"""
        
        member = {
            'nama': '',
            'provinsi': '',
            'akd': '',
            'fraksi': '',
            'foto_url': '',
            'detail_url': ''
        }
        
        try:
            # Extract photo URL from img element
            try:
                photo_img = card_element.find_element(By.CSS_SELECTOR, "img[src*='sigota/photo']")
                photo_src = photo_img.get_attribute('src')
                if photo_src:
                    member['foto_url'] = photo_src.split('?')[0] if '?' in photo_src else photo_src
            except NoSuchElementException:
                try:
                    photo_img = card_element.find_element(By.TAG_NAME, "img")
                    member['foto_url'] = photo_img.get_attribute('src')
                except:
                    pass
            
            # Extract province/dapil
            try:
                province_span = card_element.find_element(
                    By.CSS_SELECTOR, 
                    "span.flex.items-center.gap-2"
                )
                province_text = province_span.text.strip()
                member['provinsi'] = province_text
            except NoSuchElementException:
                print(f"      ⚠️  No province found for card {card_number}")
            
            # Extract name
            try:
                name_element = card_element.find_element(
                    By.CSS_SELECTOR, 
                    "p.word-break.pb-1.text-base.font-semibold"
                )
                member['nama'] = name_element.text.strip()
            except NoSuchElementException:
                print(f"      ⚠️  No name found for card {card_number}")
            
            # Extract AKD/Komisi info
            try:
                akd_list = card_element.find_elements(
                    By.CSS_SELECTOR, 
                    "ul.list-outside.list-disc li"
                )
                
                akd_items = []
                for li in akd_list:
                    akd_text = li.text.strip()
                    if akd_text:
                        akd_items.append(akd_text)
                
                member['akd'] = ', '.join(akd_items) if akd_items else ''
                
            except NoSuchElementException:
                print(f"      ⚠️  No AKD info found for card {card_number}")
            
            # Extract fraksi
            try:
                fraksi_element = card_element.find_element(
                    By.CSS_SELECTOR, 
                    "div.border-t-\\[3px\\]"
                )
                fraksi_text = fraksi_element.text.strip()
                member['fraksi'] = fraksi_text
            except NoSuchElementException:
                print(f"      ⚠️  No fraksi found for card {card_number}")
            
            # Generate detail URL
            if member['foto_url']:
                photo_match = re.search(r'/photo/(\d+)\.jpg', member['foto_url'])
                if photo_match:
                    member_id = photo_match.group(1)
                    member['detail_url'] = f"https://www.dpr.go.id/anggota/detail/id/{member_id}"
            
            return member
            
        except Exception as e:
            print(f"      ❌ Error extracting from card {card_number}: {e}")
            return member
    
    def get_total_pages(self):
        """Get total number of pages from pagination"""
        
        try:
            print("📄 Checking pagination...")
            
            pagination_container = self.driver.find_element(
                By.CSS_SELECTOR, 
                "ul.btn-group[role='navigation']"
            )
            
            page_links = pagination_container.find_elements(
                By.CSS_SELECTOR, 
                "li a[aria-label*='Page']"
            )
            
            max_page = 0
            for link in page_links:
                aria_label = link.get_attribute('aria-label')
                if aria_label and 'Page' in aria_label:
                    page_match = re.search(r'Page (\d+)', aria_label)
                    if page_match:
                        page_num = int(page_match.group(1))
                        max_page = max(max_page, page_num)
            
            if max_page > 0:
                print(f"📄 Found {max_page} total pages")
                return max_page
            else:
                print("⚠️  Could not determine total pages, defaulting to 1")
                return 1
                
        except Exception as e:
            print(f"❌ Error getting pagination: {e}")
            return 1
    
    def click_next_page(self):
        """Click next page button"""
        
        try:
            next_button = self.driver.find_element(
                By.CSS_SELECTOR, 
                "a[aria-label='Next page'][rel='next']"
            )
            
            parent_li = next_button.find_element(By.XPATH, "./..")
            if 'disabled' in parent_li.get_attribute('class'):
                print("📄 Reached last page (next button disabled)")
                return False
            
            self.driver.execute_script("arguments[0].click();", next_button)
            print("📄 Clicked next page")
            
            time.sleep(3)
            
            return True
            
        except NoSuchElementException:
            print("📄 No next button found")
            return False
        except Exception as e:
            print(f"❌ Error clicking next page: {e}")
            return False
    
    def scrape_all_pages(self):
        """Main scraping function - scrape all pages"""
        
        print("🚀 Starting DPR RI Selenium scraping...")
        print(f"📅 Scraper started at: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        
        try:
            if not self.navigate_to_page(1):
                print("❌ Failed to load first page")
                return []
            
            total_pages = self.get_total_pages()
            
            current_page = 1
            
            while current_page <= total_pages:
                print(f"\n📄 === PAGE {current_page}/{total_pages} ===")
                
                page_members = self.extract_members_from_current_page()
                
                if page_members:
                    self.members_data.extend(page_members)
                    print(f"✅ Page {current_page}: Added {len(page_members)} members")
                    print(f"📊 Total members so far: {len(self.members_data)}")
                else:
                    print(f"❌ Page {current_page}: No members found")
                
                if current_page >= total_pages:
                    print("📄 Reached final page")
                    break
                
                if current_page < total_pages:
                    if not self.click_next_page():
                        print("❌ Failed to navigate to next page")
                        break
                    
                    current_page += 1
                    time.sleep(2)
                else:
                    break
            
            print(f"\n🎯 SCRAPING COMPLETED!")
            print(f"📊 Total pages processed: {current_page}")
            print(f"👥 Total members extracted: {len(self.members_data)}")
            
            return self.members_data
            
        except KeyboardInterrupt:
            print("\n⚠️  Scraping interrupted by user")
            return self.members_data
        except Exception as e:
            print(f"❌ Unexpected error during scraping: {e}")
            return self.members_data
        finally:
            self.cleanup()
    
    def cleanup(self):
        """Clean up resources"""
        try:
            if hasattr(self, 'driver'):
                self.driver.quit()
                print("🧹 WebDriver cleaned up")
        except:
            pass
    
    def generate_simple_sentence(self, member, index):
        """Generate simple one sentence for each member optimized for RAG"""
        
        nama = member.get('nama', '').strip()
        provinsi = member.get('provinsi', '').strip()
        fraksi = member.get('fraksi', '').strip()
        akd = member.get('akd', '').strip()
        
        if not nama:
            return ""
        
        # Build one comprehensive sentence
        sentence = f"{nama} adalah anggota Dewan Perwakilan Rakyat Republik Indonesia (DPR RI)"
        
        # Add dapil/province
        if provinsi:
            sentence += f" yang mewakili daerah pemilihan {provinsi}"
        
        # Add faction
        if fraksi:
            sentence += f", merupakan anggota aktif dari {fraksi}"
        
        # Add AKD/committee info
        if akd:
            sentence += f" dan menjalankan tugas legislatif dalam {akd}"
        
        sentence += "."
        
        return sentence
    
    def create_simple_docx_report(self):
        """Create a simple DOCX report with one sentence per member"""
        
        if not self.members_data:
            print("❌ No data to create report")
            return None
        
        print("📝 Creating simple sentence DOCX report for RAG optimization...")
        
        # Create document
        doc = Document()
        
        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Add title
        title = doc.add_heading('DAFTAR ANGGOTA DPR RI', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add subtitle with timestamp
        timestamp = datetime.now().strftime('%d %B %Y, %H:%M:%S WIB')
        subtitle = doc.add_paragraph(f'Data Anggota Dewan Perwakilan Rakyat Republik Indonesia')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        date_para = doc.add_paragraph(f'Per tanggal: {timestamp}')
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add summary
        summary = doc.add_paragraph()
        summary.add_run(f"Total Anggota: {len(self.members_data)} orang").bold = True
        
        doc.add_paragraph()  # Empty line
        
        # Add each member as simple sentence
        for i, member in enumerate(self.members_data, 1):
            # Generate simple sentence
            sentence = self.generate_simple_sentence(member, i)
            
            if sentence:
                member_para = doc.add_paragraph(f"{i}. {sentence}")
                member_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        return doc
    
    def save_simple_docx_report(self):
        """Save the simple DOCX report optimized for RAG"""
        
        if not self.members_data:
            print("❌ No data to save")
            return
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"Daftar_Anggota_DPR_RI_Simple_{timestamp}.docx"
        
        try:
            doc = self.create_simple_docx_report()
            if doc:
                doc.save(filename)
                print(f"📄 Simple RAG-optimized DOCX report saved: {filename}")
                print(f"📊 Report contains {len(self.members_data)} member sentences")
                print(f"🤖 Format: One sentence per member - RAG optimized")
                self.print_summary()
            else:
                print("❌ Failed to create simple DOCX report")
        except Exception as e:
            print(f"❌ Error saving simple DOCX report: {e}")
    
    def print_summary(self):
        """Print detailed summary"""
        
        print(f"\n📊 SUMMARY:")
        print(f"=" * 50)
        print(f"📁 Output format: Simple sentences (RAG-optimized)")
        print(f"👥 Total members: {len(self.members_data)}")
        print(f"📝 Content type: One sentence per member")
        
        # Count by province
        provinces = Counter(member.get('provinsi', 'Unknown').strip() for member in self.members_data)
        fractions = Counter(member.get('fraksi', '').strip() for member in self.members_data if member.get('fraksi', '').strip())
        
        print(f"\n🗺️  Top 5 Provinces/Dapil:")
        for prov, count in provinces.most_common(5):
            print(f"   {prov}: {count} anggota")
        
        print(f"\n🏛️  Faction Distribution:")
        for frak, count in fractions.most_common():
            print(f"   {frak}: {count} anggota")
        
        # Sample sentences
        if self.members_data:
            print(f"\n📝 Sample sentences (first 3 members):")
            for i, member in enumerate(self.members_data[:3], 1):
                sentence = self.generate_simple_sentence(member, i)
                print(f"   {i}. {sentence}")
            
        print(f"\n🤖 RAG Optimization Features:")
        print(f"   ✅ One sentence per member")
        print(f"   ✅ Concise and informative")
        print(f"   ✅ Professional Indonesian language")
        print(f"   ✅ Complete essential information")

def main():
    print("🏛️  DPR RI SELENIUM SCRAPER - SIMPLE SENTENCES FOR RAG")
    print("Exports member data in simple sentence format")
    print("One sentence per member - RAG optimized")
    print("=" * 65)
    
    # Ask user preference
    print("\n🤖 Choose mode:")
    print("1. Visible browser (you can see what's happening) - Recommended")
    print("2. Headless mode (faster, no browser window)")
    
    choice = input("Enter choice (1 or 2, default=1): ").strip()
    headless = choice == "2"
    
    if headless:
        print("⚠️  Headless mode: No browser window will be shown")
    else:
        print("👀 Visible mode: Browser window will open")
    
    print("\n🤖 RAG Optimization Features:")
    print("   📝 Simple one sentence per member")
    print("   🔍 Essential information only")
    print("   🇮🇩 Clean Indonesian language")
    print("   📚 Perfect for RAG systems")
    
    scraper = DPRSeleniumSimpleSentences(headless=headless)
    
    try:
        members = scraper.scrape_all_pages()
        
        if members:
            scraper.save_simple_docx_report()
            print("\n✅ SCRAPING COMPLETED SUCCESSFULLY!")
            print(f"🎯 Total members scraped: {len(members)}")
            print("📄 Simple RAG-optimized DOCX report has been generated!")
            print("🤖 Ready for RAG processing!")
        else:
            print("\n❌ No data extracted")
            
    except KeyboardInterrupt:
        print("\n⚠️  Interrupted by user")
        if scraper.members_data:
            scraper.save_simple_docx_report()  # Save partial results
    except Exception as e:
        print(f"\n❌ Unexpected error: {e}")
        
    finally:
        try:
            scraper.cleanup()
        except:
            pass
    
    print("\n🏁 Program finished")

if __name__ == "__main__":
    main()